from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Helper: shade a table row ─────────────────────────────────────────────────
def shade_row(row, fill_hex):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill_hex)
        tcPr.append(shd)

def set_cell_border_bottom(cell, color="CCCCCC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)

# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading("AAPL Market Anomaly Detection", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph("Autoencoder-based analysis of hourly price data  |  April 2024 – April 2026")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)
sub.runs[0].font.size = Pt(10)

doc.add_paragraph()

# ── Summary box ───────────────────────────────────────────────────────────────
doc.add_heading("Summary", level=1)

summary_data = [
    ("Data",           "AAPL hourly OHLCV — 1,736 bars"),
    ("Model",          "Fully-connected autoencoder (5 → 16 → 5)"),
    ("Anomaly method", "Reconstruction MSE > mean + 3 × std"),
    ("Threshold",      "0.9278 MSE"),
    ("Anomalies found","26  (≈ 1.5 % of all bars)"),
]

tbl = doc.add_table(rows=len(summary_data), cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (k, v) in enumerate(summary_data):
    fill = "EBF3FB" if i % 2 == 0 else "FFFFFF"
    tbl.rows[i].cells[0].text = k
    tbl.rows[i].cells[1].text = v
    tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    shade_row(tbl.rows[i], fill)
    tbl.columns[0].width = Inches(2.0)
    tbl.columns[1].width = Inches(4.0)

doc.add_paragraph()

# ── Plot ──────────────────────────────────────────────────────────────────────
doc.add_heading("Anomaly Chart", level=1)
doc.add_picture("results/anomalies.png", width=Inches(6.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph("Figure 1: AAPL close price (top) and reconstruction error (bottom). Red dots = detected anomalies.")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.runs[0].font.size = Pt(9)
caption.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()

# ── Standout anomalies table ──────────────────────────────────────────────────
doc.add_heading("Standout Anomalies", level=1)

standouts = [
    ("2025-04-07", "7.72", "Largest spike — likely the tariff-driven market crash"),
    ("2025-04-14", "3.05", "Follow-on volatility after crash"),
    ("2025-09-22", "2.59", "Unusual open/close/volume pattern"),
    ("2025-09-19", "2.02", "Same volatility cluster"),
    ("2025-08-07", "1.87", "August volatility episode"),
]

headers = ["Date", "Reconstruction Error", "Likely Cause"]
tbl2 = doc.add_table(rows=1 + len(standouts), cols=3)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT

# Header row
shade_row(tbl2.rows[0], "1F497D")
for i, h in enumerate(headers):
    cell = tbl2.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# Data rows
for r, (date, err, cause) in enumerate(standouts):
    fill = "EBF3FB" if r % 2 == 0 else "FFFFFF"
    row = tbl2.rows[r + 1]
    row.cells[0].text = date
    row.cells[1].text = err
    row.cells[2].text = cause
    shade_row(row, fill)

# Column widths
for row in tbl2.rows:
    row.cells[0].width = Inches(1.3)
    row.cells[1].width = Inches(1.6)
    row.cells[2].width = Inches(3.1)

doc.add_paragraph()

# ── Full anomaly list ─────────────────────────────────────────────────────────
doc.add_heading("Full Anomaly List", level=1)

all_anomalies = [
    ("2025-04-03 13:30", "1.8372"),
    ("2025-04-04 13:30", "1.5159"),
    ("2025-04-07 13:30", "7.7223"),
    ("2025-04-09 13:30", "1.1670"),
    ("2025-04-09 18:30", "1.1667"),
    ("2025-04-14 13:30", "3.0529"),
    ("2025-05-02 13:30", "1.5079"),
    ("2025-05-23 13:30", "0.9908"),
    ("2025-06-30 18:30", "0.9743"),
    ("2025-07-01 13:30", "1.0496"),
    ("2025-08-01 13:30", "1.3840"),
    ("2025-08-07 13:30", "1.8713"),
    ("2025-08-08 16:30", "1.3767"),
    ("2025-09-03 13:30", "0.9959"),
    ("2025-09-10 13:30", "0.9495"),
    ("2025-09-16 13:30", "0.9693"),
    ("2025-09-19 13:30", "2.0206"),
    ("2025-09-22 13:30", "2.5871"),
    ("2025-09-23 18:30", "0.9723"),
    ("2025-10-20 13:30", "1.3573"),
    ("2025-10-20 14:30", "1.6512"),
    ("2025-10-21 13:30", "1.1125"),
    ("2025-10-31 13:30", "1.3918"),
    ("2025-12-19 14:30", "1.5962"),
    ("2026-01-30 14:30", "1.1136"),
    ("2026-02-04 14:30", "0.9565"),
]

tbl3 = doc.add_table(rows=1 + len(all_anomalies), cols=2)
tbl3.style = "Table Grid"
shade_row(tbl3.rows[0], "1F497D")
for i, h in enumerate(["Timestamp (UTC)", "Reconstruction Error (MSE)"]):
    cell = tbl3.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for r, (ts, err) in enumerate(all_anomalies):
    fill = "EBF3FB" if r % 2 == 0 else "FFFFFF"
    row = tbl3.rows[r + 1]
    row.cells[0].text = ts
    row.cells[1].text = err
    shade_row(row, fill)
    # Highlight the biggest outlier
    if err == "7.7223":
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        row.cells[1].paragraphs[0].runs[0].bold = True

for row in tbl3.rows:
    row.cells[0].width = Inches(2.5)
    row.cells[1].width = Inches(2.5)

doc.add_paragraph()

# ── Key finding ───────────────────────────────────────────────────────────────
doc.add_heading("Key Finding", level=1)
p = doc.add_paragraph()
p.add_run("April 7, 2025").bold = True
p.add_run(
    " is by far the most anomalous bar in the dataset, with a reconstruction error of "
)
p.add_run("7.72").bold = True
p.add_run(
    " — approximately 8× the detection threshold. This corresponds to the day AAPL "
    "dropped ~5% on tariff-related market news, producing an OHLCV pattern the model "
    "had never seen in normal conditions."
)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
footer_p = doc.add_paragraph(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}")
footer_p.runs[0].font.size = Pt(8)
footer_p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.save("results/AAPL_Anomaly_Report.docx")
print("Saved results/AAPL_Anomaly_Report.docx")
